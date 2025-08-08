### DATA PREPROSESSING ###

import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)

import pandas as pd
import re

pd.set_option('display.max_columns', None)
pd.set_option('display.width', 1500)

import os
from docx import Document

# Define the folder containing Word documents of each theme (e.g. 1-Prepare, 2-Set expectation)
folder_path = r"Paste your folder pathname here"

# Function to extract text from a Word document
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# List all Word documents in the folder
file_names = [f for f in os.listdir(folder_path) if f.endswith(".docx")]

# Read all documents and store in a list
documents = []
for file_name in file_names:
    file_path = os.path.join(folder_path, file_name)
    text_content = read_docx(file_path)
    combined_content = f"{file_name}: {text_content}"
    documents.append({"Combined": combined_content})

# Convert the list to a DataFrame
xt = pd.DataFrame(documents)
print(xt)

def clean_input_data(text):
    # Remove leading/trailing spaces
    text = text.strip()
    # Remove special characters (optional, based on requirements)
    text = text.replace('\n', ' ').replace('\r', '')
    text = re.sub(r'(?<=\w)-(?=\w)', '', text)
    # text = re.sub('variable', '', text)
    text = re.sub('remotely', 'remote', text)  # Change words to what you need to be removed
    return text

xt['Combined'] = xt['Combined'].apply(clean_input_data)
xt['Combined'] = xt['Combined'].astype(str)
xt.info()


import nltk
import matplotlib.pyplot as plt
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from nltk.corpus import stopwords
from nltk import pos_tag
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet
from nltk.util import ngrams
from collections import Counter

# Step 1: Preprocess the Text Data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('punkt_tab')
nltk.download('averaged_perceptron_tagger')
nltk.download('popular')

import nltk
nltk.download('averaged_perceptron_tagger')

# Initialize the lemmatizer
lemmatizer = WordNetLemmatizer()

from nltk import pos_tag

def get_wordnet_pos(treebank_tag):
    """Convert POS tag from nltk's pos_tag to wordnet POS tag."""
    if treebank_tag.startswith('J'):
        return wordnet.ADJ
    elif treebank_tag.startswith('V'):
        return wordnet.VERB
    elif treebank_tag.startswith('N'):
        return wordnet.NOUN
    elif treebank_tag.startswith('R'):
        return wordnet.ADV
    else:
        return wordnet.NOUN  # Default to noun if the tag is not recognized

def preprocess_text(text):
    # Convert text to lowercase
    text = text.lower()
    # Tokenize the text into words
    tokens = nltk.word_tokenize(text)
    # Remove punctuation and non-alphabetic tokens
    tokens = [word for word in tokens if word.isalpha() or '-' in word]
    # Remove stopwords (works with text in English)
    stop_words = set(stopwords.words('english'))
    stop_words.update(['like', 'would', 'yeah', 'uh', 'hey'])  # Add words you need to be removed but StopWord didn't remove
    tokens = [word for word in tokens if word not in stop_words]
    # POS tagging
    tagged_tokens = pos_tag(tokens)
    # Lemmatize each token with its POS tag
    tokens = [lemmatizer.lemmatize(word, get_wordnet_pos(tag)) for word, tag in tagged_tokens]
    return ' '.join(tokens)  # return tokens and Convert list to string

# Apply the preprocessing function
xt['processed'] = xt['Combined'].apply(lambda x: preprocess_text(x))

# Display processed text
print(xt[['processed']].head())

import os
import docx
import pandas as pd
import matplotlib.pyplot as plt
from collections import Counter

# Load text from all .docx files in a given folder
def load_docx_text(folder_path):
    all_text = ""
    for file in os.listdir(folder_path):
        if file.endswith(".docx"):
            doc = docx.Document(os.path.join(folder_path, file))
            for para in doc.paragraphs:
                all_text += para.text.lower() + " "
    return all_text


# Count occurrences of words in a given text
def count_word_occurrences(word_list, text):
    text_words = text.split()  # Simple tokenization (can be improved)
    total_words = len(text_words)
    word_counts = Counter(text_words)

    row_counts = {word: word_counts[ word ] for word in word_list}
    row_total = sum(row_counts.values())
    proportion = (row_total / total_words) * 100 if total_words > 0 else 0

    return row_total, total_words, proportion


# Visualize the results
def plot_proportions(proportions):
    labels = [ f"Row {i}" for i in range(len(proportions)) ]
    values = [ prop for prop in proportions ]

    plt.figure(figsize=(8, 6))
    plt.bar(labels, values, color=[ 'blue', 'green', 'red', 'purple', 'orange' ])
    plt.xlabel("Rows")
    plt.ylabel("Proportion (%)")
    plt.title("Word Proportion in Document Text")
    plt.ylim(0, 100)

    for i, v in enumerate(values):
        plt.text(i, v + 1, f"{v:.2f}%", ha='center', fontsize=12)

    plt.show()


# Change folder path to the one where your raw transcript is
# If you have several transcripts add one to the folder at a time and
# run only the code below this section again to get the percentage.

folder_path_1 = r'Add your new folder Pathname here'
# Load and process text
doc_text = load_docx_text(folder_path_1)
proportions = [ ]

for i, row in xt.iterrows():
    words = row['processed'].split()
    row_total, total_words, proportion = count_word_occurrences(words, doc_text)
    proportions.append(proportion)

print(proportions)
